import os
import json
import threading
from datetime import datetime

from flask import (
    Flask,
    request,
    jsonify,
    render_template,
    send_file,
    abort,
)

from recorder import Recorder, RecorderError
from transcribe import transcribe, TranscribeError
from analyze import analyze, AnalyzeError

BASE = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__, template_folder=os.path.join(BASE, "templates"))

# ---- состояние (в памяти) ----
CONFIG = {
    "stt_provider": "yandex",  # yandex | openai (Whisper) | local (faster-whisper)
    "stt_api_key": "",
    "stt_folder_id": "",
    "stt_model": "whisper-1",
    "local_model": "small",
    "analysis_provider": "local",  # local (llama.cpp) | openai
    "llm_model": "qwen2.5-1.5b-instruct-q4_k_m.gguf",
    "openai_api_key": "",
    "openai_model": "gpt-4o-mini",
    "record_mode": "auto",
}

recorder = None
STATE = {
    "status": "idle",  # idle | recording | processing | done | error
    "message": "",
    "progress": 0.0,
    "transcript": "",
    "report": None,
    "report_docx_path": None,
}


def _set_status(status, message="", progress=None):
    STATE["status"] = status
    STATE["message"] = message
    if progress is not None:
        STATE["progress"] = progress


@app.route("/")
def index():
    return render_template("index.html", config=CONFIG, state=STATE)


@app.route("/api/config", methods=["POST"])
def set_config():
    data = request.get_json(force=True, silent=True) or {}
    for k in CONFIG:
        if k in data:
            CONFIG[k] = data[k]
    return jsonify({"ok": True, "config": CONFIG})


@app.route("/api/start", methods=["POST"])
def start_rec():
    global recorder
    if recorder and recorder.is_running():
        return jsonify({"ok": False, "error": "Запись уже идёт."})
    try:
        recorder = Recorder(mode=CONFIG.get("record_mode") or "auto")
        recorder.start()
    except RecorderError as e:
        return jsonify({"ok": False, "error": str(e)})
    _set_status("recording", "Идёт запись встречи...")
    return jsonify({"ok": True})


@app.route("/api/stop", methods=["POST"])
def stop_rec():
    global recorder
    if not recorder or not recorder.is_running():
        return jsonify({"ok": False, "error": "Запись не активна."})

    _set_status("processing", "Остановка записи и сохранение WAV...")
    wav_path = recorder.stop()
    recorder = None

    if not wav_path:
        _set_status("error", "Не удалось получить аудио (пустая запись).")
        return jsonify({"ok": False, "error": STATE["message"]})

    try:
        _set_status("processing", "Распознавание речи...")
        transcript = transcribe(
            wav_path,
            CONFIG.get("stt_provider", "yandex"),
            progress_cb=lambda p, _t: _set_status(
                "processing", "Распознавание речи...", p
            ),
            api_key=CONFIG["stt_api_key"],
            folder_id=CONFIG["stt_folder_id"],
            openai_key=CONFIG["openai_api_key"],
            openai_model=CONFIG.get("stt_model", "whisper-1"),
            local_model=CONFIG.get("local_model", "small"),
        )
    except TranscribeError as e:
        _set_status("error", f"Ошибка STT: {e}")
        return jsonify({"ok": False, "error": str(e)})

    if not transcript.strip():
        _set_status("done", "Речь не распознана.", 1.0)
        STATE["transcript"] = ""
        STATE["report"] = _empty_report()
        return jsonify({"ok": True, "transcript": "", "report": STATE["report"]})

    try:
        _set_status("processing", "Анализ и извлечение задач (локальная LLM)...")
        report = analyze(
            transcript,
            CONFIG["analysis_provider"],
            progress_cb=lambda p, _t: _set_status(
                "processing", "Анализ встречи (LLM)...", 0.5 + 0.5 * p
            ),
            llm_model=CONFIG["llm_model"],
            openai_key=CONFIG["openai_api_key"],
            openai_model=CONFIG["openai_model"],
        )
    except AnalyzeError as e:
        _set_status("error", f"Ошибка анализа: {e}")
        return jsonify({"ok": False, "error": str(e)})

    report.setdefault("summary", "")
    report.setdefault("tasks", [])
    report.setdefault("deadlines", [])
    report.setdefault("decisions", [])

    STATE["transcript"] = transcript
    STATE["report"] = report
    STATE["report_docx_path"] = _write_report_docx(report, transcript)

    _set_status("done", "Готово!", 1.0)
    return jsonify({"ok": True, "transcript": transcript, "report": report})


def _empty_report():
    return {"summary": "", "tasks": [], "deadlines": [], "decisions": []}


def _norm_decision(d):
    if isinstance(d, dict):
        return d.get("description") or d.get("decision") or " ".join(str(v) for v in d.values())
    return d


def _write_report_docx(report, transcript):
    from docx import Document

    doc = Document()
    doc.add_heading(f"Итог встречи — {datetime.now():%Y-%m-%d %H:%M}", level=1)

    doc.add_heading("Краткий итог", level=2)
    doc.add_paragraph(report.get("summary") or "_(пусто)_")

    doc.add_heading("Задачи", level=2)
    tasks = report.get("tasks") or []
    if tasks:
        for i, t in enumerate(tasks, 1):
            owner = (t.get("owner") if isinstance(t, dict) else None) or "—"
            dl = (t.get("deadline") if isinstance(t, dict) else None) or "—"
            dl_iso = (t.get("deadline_iso") if isinstance(t, dict) else None) or None
            dl_str = f"{dl} ({dl_iso})" if dl_iso else dl
            desc = (t.get("description") if isinstance(t, dict) else None) or ""
            doc.add_paragraph(
                f"{i}. {desc}  (Ответственный: {owner}; Срок: {dl_str})", style="List Number"
            )
    else:
        doc.add_paragraph("_(задач не выявлено)_")

    doc.add_heading("Сроки", level=2)
    dls = report.get("deadlines") or []
    if dls:
        for d in dls:
            what = (d.get("what") if isinstance(d, dict) else None) or ""
            when = (d.get("when") if isinstance(d, dict) else None) or ""
            iso = (d.get("when_iso") if isinstance(d, dict) else None) or None
            when_str = f"{when} ({iso})" if iso else when
            doc.add_paragraph(f"{what}: {when_str}", style="List Bullet")
    else:
        doc.add_paragraph("_(сроков не выявлено)_")

    doc.add_heading("Принятые решения", level=2)
    decs = report.get("decisions") or []
    if decs:
        for d in decs:
            doc.add_paragraph(_norm_decision(d), style="List Bullet")
    else:
        doc.add_paragraph("_(решений не выявлено)_")

    doc.add_heading("Полная расшифровка", level=2)
    doc.add_paragraph(transcript or "")

    doc_date = datetime.now().strftime("%d-%m-%Y")
    path = os.path.join(BASE, f"meeting_report_{doc_date}.docx")
    doc.save(path)
    return path


@app.route("/api/status")
def status():
    return jsonify(STATE)


@app.route("/api/report/docx")
def report_docx():
    p = STATE.get("report_docx_path")
    if not p or not os.path.exists(p):
        abort(404)
    doc_date = datetime.now().strftime("%d-%m-%Y")
    return send_file(
        p,
        as_attachment=True,
        download_name=f"meeting_report_{doc_date}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    print(f"Running on http://127.0.0.1:{port}", flush=True)
    app.run(host="127.0.0.1", port=port, debug=False)
